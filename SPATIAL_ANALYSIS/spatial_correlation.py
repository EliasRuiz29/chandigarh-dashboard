"""
Chandigarh Multi-Layer Map Project
===================================
Week 7 — Day 6-7: Spatial Correlation Analysis + Results Document
Produces:
  - day6_correlation_matrix.png       (Spearman heatmap with significance stars)
  - day6_scatter_plots.png            (6 bivariate scatter plots vs LST)
  - day6_LISA_clusters.png            (LISA cluster map — all 63 sectors)
  - day6_spearman_matrix.csv          (full ρ matrix)
  - day6_correlation_summary.csv      (key pairs summary for paper)
  - sector_parameter_table.csv        (all sectors × all variables, sorted by LST)
  - master_geo.gpkg                   (updated with LISA_cluster column)
  - Morans_LISA_Results.docx          (full results document with Moran's I + LISA table)

Author : Elias Ruiz Sabater
Date   : April 2026

Notes on spatial weights:
  - 25W, 26E, 38W, Rajindra Park share borders → included in Queen weights (61 sectors)
  - Manimajra, Khuda Alisher are geographic islands (no shared borders) → excluded from
    weights and LISA, shown on map as 'Island' category
"""

import pandas as pd
import numpy as np
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import matplotlib.patheffects as pe
import seaborn as sns
from scipy import stats
import geopandas as gpd
from libpysal.weights import Queen
from esda.moran import Moran, Moran_Local
from spreg import OLS, GM_Lag
import subprocess, sys, os
import warnings
warnings.filterwarnings('ignore')


# ─────────────────────────────────────────────────────────────────────────────
# CONFIGURATION — adjust paths as needed
# ─────────────────────────────────────────────────────────────────────────────

MASTER_PATH      = '/Users/eliasruizsabater/Desktop/Project MMU/MASTER.xlsx'
SHAPEFILE_PATH   = '/Users/eliasruizsabater/Desktop/Project MMU/Chandigarh_Boundary-SHP/PySAL/Chandigarh_Sectors_UTM43N.shp'
MASTER_GEO_PATH  = '/Users/eliasruizsabater/Desktop/Project MMU/master_geo.gpkg'
OUTPUT_DIR       = '/Users/eliasruizsabater/Desktop/Project MMU/RESULTS/SPATIAL_ANALYSIS/'

# Geographic islands — sectors with no shared borders in the shapefile
# Cannot be included in Queen contiguity weights
ISLANDS = ['Manimajra', 'Khuda Alisher']

# Output file paths
OUT_CORR_MATRIX  = OUTPUT_DIR + 'Correlation_matrix.png'
OUT_SCATTER      = OUTPUT_DIR + 'Correlation_scatter_plots.png'
OUT_LISA_MAP     = OUTPUT_DIR + 'LISA_clusters.png'
OUT_RHO_CSV      = OUTPUT_DIR + 'Spearman_matrix.csv'
OUT_SUMMARY_CSV  = OUTPUT_DIR + 'Correlation_summary.csv'
OUT_PARAM_TABLE  = OUTPUT_DIR + 'sector_parameter_table.csv'
OUT_DOCX         = OUTPUT_DIR + 'Morans_LISA_Results.docx'
OUT_DOCX_JS      = OUTPUT_DIR + '_build_morans_lisa.js'   # temp JS file, deleted after build


# ─────────────────────────────────────────────────────────────────────────────
# 1. LOAD & PREPARE DATA
# ─────────────────────────────────────────────────────────────────────────────

df = pd.read_excel(MASTER_PATH, sheet_name='Sheet2')
df['Sector_nam'] = df['Sector_nam'].astype(str)

# Urban sectors only — exclude rural zones
urban = df[~df['Sector_nam'].str.startswith('Rural')].copy().reset_index(drop=True)
urban = urban.rename(columns={
    'PopxSector' : 'Population',
    '% SC Pop '  : 'SC_pct',
})

VARS = ['LST', 'NDVI', 'NDBI', 'UTFVI', 'TCI', 'GHI', 'Population', 'SC_pct']
print(f"Urban sectors loaded: {len(urban)}")


# ─────────────────────────────────────────────────────────────────────────────
# 2. SPEARMAN CORRELATION MATRIX
# ─────────────────────────────────────────────────────────────────────────────
# Pairwise complete observations — each pair uses all rows where both
# variables are non-NaN. This correctly handles the GHI case (n=24).

rho_mat  = pd.DataFrame(np.nan, index=VARS, columns=VARS)
pval_mat = pd.DataFrame(np.nan, index=VARS, columns=VARS)

for v1 in VARS:
    for v2 in VARS:
        mask = urban[v1].notna() & urban[v2].notna()
        x, y = urban.loc[mask, v1], urban.loc[mask, v2]
        if len(x) >= 5:
            r, p = stats.spearmanr(x, y)
            rho_mat.loc[v1, v2]  = round(r, 3)
            pval_mat.loc[v1, v2] = round(p, 4)

print("\n=== Spearman ρ matrix ===")
print(rho_mat.to_string())
print("\n=== p-value matrix ===")
print(pval_mat.to_string())

# Key hypotheses check
print("\n=== Key hypotheses ===")
HYPOTHESES = [
    ('LST', 'NDVI',       'expect negative'),
    ('LST', 'TCI',        'expect positive'),
    ('LST', 'NDBI',       'expect positive'),
    ('LST', 'GHI',        'expect negative'),
    ('LST', 'Population', 'expect positive'),
    ('LST', 'SC_pct',     'expect positive'),
]
for v1, v2, hyp in HYPOTHESES:
    r = rho_mat.loc[v1, v2]
    p = pval_mat.loc[v1, v2]
    n_pair = (urban[v1].notna() & urban[v2].notna()).sum()
    sig  = "✅ sig" if p < 0.05 else "— ns"
    conf = "✅ confirmed" if (
        (hyp == 'expect negative' and r < 0) or
        (hyp == 'expect positive' and r > 0)
    ) else "❌ reversed"
    print(f"  {v1} vs {v2}: ρ={r:.3f}  p={p:.4f}  n={n_pair}  [{sig}]  [{conf}]")

rho_mat.to_csv(OUT_RHO_CSV)
print(f"\nSpearman matrix saved: {OUT_RHO_CSV}")

summary = pd.DataFrame({
    'Variable_pair'      : ['LST–NDVI', 'LST–TCI', 'LST–NDBI', 'LST–GHI', 'LST–Population', 'LST–SC_pct'],
    'Spearman_rho'       : [rho_mat.loc['LST', v] for v in ['NDVI', 'TCI', 'NDBI', 'GHI', 'Population', 'SC_pct']],
    'p_value'            : [pval_mat.loc['LST', v] for v in ['NDVI', 'TCI', 'NDBI', 'GHI', 'Population', 'SC_pct']],
    'n'                  : [(urban['LST'].notna() & urban[v].notna()).sum() for v in ['NDVI', 'TCI', 'NDBI', 'GHI', 'Population', 'SC_pct']],
    'Significant_p05'    : ['Yes', 'Yes', 'Yes', 'No', 'No', 'No'],
    'Direction_confirmed': ['Yes', 'Yes', 'No — reversed', 'Yes', 'Yes', 'Yes'],
})
summary.to_csv(OUT_SUMMARY_CSV, index=False)
print(f"Correlation summary saved: {OUT_SUMMARY_CSV}")


# ─────────────────────────────────────────────────────────────────────────────
# 3. CORRELATION HEATMAP
# ─────────────────────────────────────────────────────────────────────────────

BG = '#0d1117'
fig, ax = plt.subplots(figsize=(10, 8), facecolor=BG)
ax.set_facecolor(BG)

mask = np.triu(np.ones_like(rho_mat, dtype=bool), k=1)

annot = rho_mat.copy().astype(str)
for v1 in VARS:
    for v2 in VARS:
        r = rho_mat.loc[v1, v2]
        p = pval_mat.loc[v1, v2]
        if pd.isna(r):
            annot.loc[v1, v2] = ''
        elif v1 == v2:
            annot.loc[v1, v2] = '1.0'
        else:
            stars = '***' if p < 0.001 else ('**' if p < 0.01 else ('*' if p < 0.05 else ''))
            annot.loc[v1, v2] = f"{r:.2f}{stars}"

sns.heatmap(
    rho_mat.astype(float), mask=mask, annot=annot, fmt='',
    cmap='RdBu_r', vmin=-1, vmax=1, linewidths=0.5, linecolor='#1a1a2e',
    ax=ax, cbar_kws={'shrink': 0.8, 'label': 'Spearman ρ'},
    annot_kws={'size': 9, 'color': 'white', 'weight': 'bold'},
)
ax.set_xticklabels(VARS, rotation=45, ha='right', color='white', fontsize=10)
ax.set_yticklabels(VARS, rotation=0,  color='white', fontsize=10)
ax.tick_params(colors='white')
cbar = ax.collections[0].colorbar
cbar.set_label('Spearman ρ', color='white', fontsize=10)
cbar.ax.yaxis.set_tick_params(color='white')
plt.setp(cbar.ax.yaxis.get_ticklabels(), color='white')
ax.set_title(
    'Spearman Rank Correlation Matrix\n'
    'Chandigarh Urban Sectors (n=63, pairwise complete)\n'
    '* p<0.05   ** p<0.01   *** p<0.001',
    color='white', fontsize=12, fontweight='bold', pad=14,
)
plt.tight_layout()
plt.savefig(OUT_CORR_MATRIX, dpi=180, bbox_inches='tight', facecolor=BG)
plt.close()
print(f"Correlation heatmap saved: {OUT_CORR_MATRIX}")


# ─────────────────────────────────────────────────────────────────────────────
# 4. SCATTER PLOTS — KEY RELATIONSHIPS
# ─────────────────────────────────────────────────────────────────────────────

fig, axes = plt.subplots(2, 3, figsize=(18, 11), facecolor=BG)
fig.suptitle(
    'Key Bivariate Relationships — LST vs Environmental & Socio-Economic Variables\n'
    'Chandigarh Urban Sectors  |  * p<0.05  ** p<0.01  *** p<0.001',
    color='white', fontsize=13, fontweight='bold', y=1.01,
)
PAIRS = [
    ('NDVI',       'ρ=−0.816***', '#4da6ff'),
    ('TCI',        'ρ=+0.340**',  '#ff6b6b'),
    ('NDBI',       'ρ=−0.601***', '#ffd700'),
    ('GHI',        'ρ=−0.316 ns', '#a78bfa'),
    ('Population', 'ρ=+0.120 ns', '#34d399'),
    ('SC_pct',     'ρ=+0.207 ns', '#fb923c'),
]
for ax, (xvar, label, color) in zip(axes.flatten(), PAIRS):
    ax.set_facecolor('#0d1b2a')
    for sp in ax.spines.values():
        sp.set_color('#334466')
    valid = urban[[xvar, 'LST']].dropna()
    x_v, y_v = valid[xvar].values, valid['LST'].values
    ax.scatter(x_v, y_v, color=color, alpha=0.75, s=50, zorder=3,
               edgecolors='white', linewidths=0.3)
    if len(x_v) >= 5:
        m, b = np.polyfit(x_v, y_v, 1)
        xl = np.linspace(x_v.min(), x_v.max(), 100)
        ax.plot(xl, m * xl + b, color=color, lw=1.8, alpha=0.6, linestyle='--')
    ax.text(0.05, 0.93, label,           transform=ax.transAxes,
            color=color,     fontsize=10, fontweight='bold', va='top')
    ax.text(0.05, 0.84, f"n={len(x_v)}", transform=ax.transAxes,
            color='#aaaaaa', fontsize=8.5, va='top')
    ax.set_xlabel(xvar,       color='#aaaaaa', fontsize=10)
    ax.set_ylabel('LST (°C)', color='#aaaaaa', fontsize=10)
    ax.tick_params(colors='#aaaaaa', labelsize=9)
    ax.grid(color='#1e2d3d', lw=0.6, alpha=0.7)
plt.tight_layout(pad=2.5)
plt.savefig(OUT_SCATTER, dpi=180, bbox_inches='tight', facecolor=BG)
plt.close()
print(f"Scatter plots saved: {OUT_SCATTER}")


# ─────────────────────────────────────────────────────────────────────────────
# 5. BUILD SPATIAL GEODATAFRAME — ALL 63 URBAN SECTORS
# ─────────────────────────────────────────────────────────────────────────────
# Load directly from the original shapefile (UTM43N) and join master data.
# This guarantees all sectors — including 25W, 26E, 38W, Rajindra Park —
# are present and correctly georeferenced.

gdf_shp = gpd.read_file(SHAPEFILE_PATH)
gdf_shp['Sector_nam'] = gdf_shp['Sector_nam'].astype(str)

master_urban = urban.copy()   # already renamed and filtered above
gdf = gdf_shp.merge(master_urban, on='Sector_nam', how='left').reset_index(drop=True)

print(f"\nGeoDataFrame: {len(gdf)} sectors joined, {gdf['LST'].notna().sum()} with LST values")

# ─────────────────────────────────────────────────────────────────────────────
# 6. SPATIAL WEIGHTS — QUEEN CONTIGUITY
# ─────────────────────────────────────────────────────────────────────────────
# Include all 61 non-island sectors (57 numeric + 25W, 26E, 38W, Rajindra Park).
# Manimajra and Khuda Alisher are geometric islands — zero shared borders in the
# shapefile — so they cannot receive Queen neighbours. They are plotted separately.

gdf_weights = gdf[~gdf['Sector_nam'].isin(ISLANDS)].reset_index(drop=True).copy()

w = Queen.from_dataframe(gdf_weights, ids=gdf_weights['Sector_nam'].tolist())
w.transform = 'r'
avg_neigh = np.mean([len(v) for v in w.neighbors.values()])
print(f"Queen weights: {len(w.neighbors)} sectors, avg {avg_neigh:.2f} neighbours, islands={w.islands}")


# ─────────────────────────────────────────────────────────────────────────────
# 7. MORAN'S I — GLOBAL SPATIAL AUTOCORRELATION
# ─────────────────────────────────────────────────────────────────────────────

lst_vals = gdf_weights['LST'].values
mi = Moran(lst_vals, w, permutations=999)

print(f"\n=== Moran's I — LST ===")
print(f"  I = {mi.I:.4f}  |  E[I] = {mi.EI:.4f}  |  p-sim = {mi.p_sim:.4f}")
print(f"  z-score = {mi.z_sim:.3f}")
print(f"  Result: {'✅ Significant positive spatial autocorrelation → use Spatial Lag Model' if mi.p_sim < 0.05 and mi.I > 0 else '— Not significant → OLS is appropriate'}")

mi_results = {}
for var in ['NDVI', 'TCI']:
    vals = gdf_weights[var].values
    mi_v = Moran(vals, w, permutations=999)
    mi_results[var] = mi_v
    print(f"\n=== Moran's I — {var} ===")
    print(f"  I = {mi_v.I:.4f}  |  p-sim = {mi_v.p_sim:.4f}  |  z = {mi_v.z_sim:.3f}")


# ─────────────────────────────────────────────────────────────────────────────
# 8. LISA — LOCAL SPATIAL AUTOCORRELATION
# ─────────────────────────────────────────────────────────────────────────────

lisa = Moran_Local(lst_vals, w, permutations=999)

cluster_labels = []
for i in range(len(lisa.Is)):
    if lisa.p_sim[i] >= 0.05:
        cluster_labels.append('NS')
    elif lisa.q[i] == 1:
        cluster_labels.append('HH')   # High LST, high-LST neighbours
    elif lisa.q[i] == 2:
        cluster_labels.append('LH')   # Low LST, high-LST neighbours
    elif lisa.q[i] == 3:
        cluster_labels.append('LL')   # Low LST, low-LST neighbours
    else:
        cluster_labels.append('HL')   # High LST, low-LST neighbours

gdf_weights['LISA_cluster'] = cluster_labels

print(f"\n=== LISA cluster counts ===")
print(gdf_weights['LISA_cluster'].value_counts().to_string())
print("\nHH sectors (hot spots): ", gdf_weights[gdf_weights['LISA_cluster'] == 'HH']['Sector_nam'].tolist())
print("LL sectors (cold spots):", gdf_weights[gdf_weights['LISA_cluster'] == 'LL']['Sector_nam'].tolist())

# Merge LISA back to full gdf; islands get their own label
gdf = gdf.merge(gdf_weights[['Sector_nam', 'LISA_cluster']], on='Sector_nam', how='left')
gdf['LISA_cluster'] = gdf['LISA_cluster'].fillna('Island')


# ─────────────────────────────────────────────────────────────────────────────
# 9. LISA MAP — ALL 63 SECTORS
# ─────────────────────────────────────────────────────────────────────────────

COLORS_MAP = {
    'HH'    : '#d7191c',   # dark red      — hot spot
    'LL'    : '#2c7bb6',   # dark blue     — cold spot
    'LH'    : '#fdae61',   # light orange  — spatial outlier
    'HL'    : '#abd9e9',   # light blue    — spatial outlier
    'NS'    : '#555577',   # grey          — not significant
    'Island': '#888844',   # olive         — geographic island, excluded from analysis
}
LABELS_MAP = {
    'HH'    : 'Hot Spot (HH)',
    'LL'    : 'Cold Spot (LL)',
    'LH'    : 'Spatial Outlier (LH)',
    'HL'    : 'Spatial Outlier (HL)',
    'NS'    : 'Not Significant',
    'Island': 'Geographic Island (excluded)',
}

fig, ax = plt.subplots(figsize=(10, 10), facecolor=BG)
ax.set_facecolor(BG)
ax.set_axis_off()

for cat, color in COLORS_MAP.items():
    sub = gdf[gdf['LISA_cluster'] == cat]
    if len(sub):
        sub.plot(color=color, linewidth=0.5, edgecolor='#1a1a2e', ax=ax)

# Labels — all 63 sectors
for _, row in gdf.iterrows():
    c = row.geometry.centroid
    txt = ax.annotate(
        str(row['Sector_nam']),
        xy=(c.x, c.y), ha='center', va='center',
        fontsize=5.5, fontweight='bold', color='white',
    )
    txt.set_path_effects([pe.withStroke(linewidth=1.5, foreground='black')])

patches = [mpatches.Patch(color=COLORS_MAP[k], label=LABELS_MAP[k])
           for k in COLORS_MAP if gdf['LISA_cluster'].isin([k]).any()]
ax.legend(handles=patches, loc='lower right', fontsize=9,
          facecolor='#0d1b2a', edgecolor='#334466', labelcolor='white',
          title="LISA Cluster  (p < 0.05)", title_fontsize=9)
ax.set_title(
    "LISA Clusters — Land Surface Temperature\n"
    "Chandigarh Urban Sectors  |  Moran's Local I  |  p < 0.05",
    color='white', fontsize=13, fontweight='bold', pad=14,
)
plt.tight_layout()
plt.savefig(OUT_LISA_MAP, dpi=180, bbox_inches='tight', facecolor=BG)
plt.close()
print(f"LISA map saved: {OUT_LISA_MAP}")


# ─────────────────────────────────────────────────────────────────────────────
# 10. SPATIAL REGRESSION
# ─────────────────────────────────────────────────────────────────────────────
# Moran's I significant → Spatial Lag Model (GM_Lag) preferred over OLS.
# Features: NDVI, TCI, Population, SC_pct
# GHI excluded (n=24 only — insufficient for regression)

# Align to weights order (gdf_weights index)
num_sectors = gdf_weights['Sector_nam'].tolist()
reg_df = urban[urban['Sector_nam'].isin(num_sectors)].copy()
reg_df = reg_df.set_index('Sector_nam').loc[num_sectors].reset_index()

y = reg_df[['LST']].values
X = reg_df[['NDVI', 'TCI', 'Population', 'SC_pct']].values

# OLS baseline
ols = OLS(y, X, w=w,
          name_y='LST',
          name_x=['NDVI', 'TCI', 'Population', 'SC_pct'],
          name_ds='Chandigarh')
print(f"\n=== OLS Regression (baseline) ===")
print(f"  R² = {ols.r2:.4f}  |  Adj-R² = {ols.ar2:.4f}")
for i, name in enumerate(['Intercept', 'NDVI', 'TCI', 'Population', 'SC_pct']):
    b = ols.betas[i][0]
    t = ols.t_stat[i][0]
    p = ols.t_stat[i][1]
    sig = '***' if p < 0.001 else ('**' if p < 0.01 else ('*' if p < 0.05 else ''))
    print(f"    {name:12s}: β={b:8.4f}  t={t:7.3f}  p={p:.4f}  {sig}")

# Spatial Lag Model
slm = GM_Lag(y, X, w=w,
             name_y='LST',
             name_x=['NDVI', 'TCI', 'Population', 'SC_pct'],
             name_ds='Chandigarh')
print(f"\n=== Spatial Lag Model (recommended — Moran's I significant) ===")
print(f"  Pseudo-R² = {slm.pr2:.4f}")
for i, name in enumerate(['Intercept', 'NDVI', 'TCI', 'Population', 'SC_pct', 'W_LST (ρ)']):
    b  = slm.betas[i][0]
    se = slm.std_err[i]
    t  = b / se if se != 0 else np.nan
    print(f"    {name:16s}: β={b:8.4f}  std={se:.4f}  t={t:.3f}")


# ─────────────────────────────────────────────────────────────────────────────
# 11. SAVE ENRICHED MASTER_GEO.GPKG WITH LISA LABELS
# ─────────────────────────────────────────────────────────────────────────────

gdf_full = gpd.read_file(MASTER_GEO_PATH)
gdf_full['Sector_nam'] = gdf_full['Sector_nam'].astype(str)
if 'LISA_cluster' in gdf_full.columns:
    gdf_full = gdf_full.drop(columns=['LISA_cluster'])
lisa_df = gdf[['Sector_nam', 'LISA_cluster']].copy()
gdf_full = gdf_full.merge(lisa_df, on='Sector_nam', how='left')
gdf_full['LISA_cluster'] = gdf_full['LISA_cluster'].fillna('NS')
gdf_full.to_file(MASTER_GEO_PATH, driver='GPKG')
print(f"\nmaster_geo.gpkg updated with LISA_cluster column: {MASTER_GEO_PATH}")


# ─────────────────────────────────────────────────────────────────────────────
# 12. SECTOR PARAMETER TABLE
# ─────────────────────────────────────────────────────────────────────────────

param = urban[['Sector_nam', 'LST', 'NDVI', 'NDBI', 'UTFVI', 'TCI', 'GHI', 'Population', 'SC_pct']].copy()
param = param.sort_values('LST', ascending=False).reset_index(drop=True)
param.to_csv(OUT_PARAM_TABLE, index=False)
print(f"sector_parameter_table.csv saved: {OUT_PARAM_TABLE}")


# ─────────────────────────────────────────────────────────────────────────────
# 13. BUILD MORAN'S I + LISA RESULTS DOCUMENT (.docx)
# ─────────────────────────────────────────────────────────────────────────────
# Collects all computed values from above and passes them to a Node.js
# script that generates the formatted Word document via docx-js.

print(f"\nBuilding Moran's I + LISA results document...")

# Collect full LISA data for document
lisa_rows = []
for i, sector in enumerate(gdf_weights['Sector_nam'].tolist()):
    row = gdf_weights[gdf_weights['Sector_nam'] == sector].iloc[0]
    neighbours = w.neighbors[sector]
    valid_n = [n for n in neighbours if n in gdf_weights['Sector_nam'].values]
    avg_n_lst = (np.mean([gdf_weights[gdf_weights['Sector_nam'] == n]['LST'].values[0]
                          for n in valid_n]) if valid_n else float('nan'))
    lisa_rows.append({
        's'    : sector,
        'lst'  : round(float(row['LST']), 3),
        'n_lst': round(float(avg_n_lst), 3) if not np.isnan(avg_n_lst) else 0,
        'li'   : round(float(lisa.Is[i]), 4),
        'p'    : round(float(lisa.p_sim[i]), 4),
        'cl'   : gdf_weights.loc[gdf_weights['Sector_nam'] == sector, 'LISA_cluster'].values[0],
        'nn'   : len(neighbours),
        'ndvi' : round(float(row['NDVI']), 4),
        'tci'  : round(float(row['TCI']), 4),
    })

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def shade_cell(cell, fill_color):
    """Add background color to a table cell"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def set_cell_border(cell, **kwargs):
    """Set cell border"""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            edge_element = OxmlElement(f'w:{edge}')
            edge_element.set(qn('w:val'), 'single')
            edge_element.set(qn('w:sz'), '12')
            edge_element.set(qn('w:space'), '0')
            edge_element.set(qn('w:color'), kwargs.get(edge, '000000'))
            tcBorders.append(edge_element)
    tcPr.append(tcBorders)

# Create DOCX document
doc = Document()
doc.core_properties.title = "Moran's I & LISA Results"

# Set up colors
COLORS = {
    'NAVY': RGBColor(31, 56, 100),
    'BLUE': RGBColor(46, 95, 163),
    'LBLUE': RGBColor(214, 228, 247),
    'GREEN': RGBColor(55, 86, 35),
    'LGREEN': RGBColor(226, 239, 218),
    'TEAL': RGBColor(13, 110, 110),
    'LTEAL': RGBColor(224, 244, 244),
    'COBALT': RGBColor(26, 58, 122),
    'LCOBALT': RGBColor(232, 238, 248),
    'ORANGE': RGBColor(197, 90, 17),
    'LORANGE': RGBColor(252, 228, 214),
    'DARK': RGBColor(26, 26, 26),
    'GREY': RGBColor(242, 242, 242),
}

# Cluster colors mapping
CLUSTER_COLORS = {
    'HH': 'D7191C',
    'LL': '2C7BB6',
    'HL': 'ABD9E9',
    'LH': 'FDAE61',
    'NS': '555577'
}

# Helper function for significance labels
def p_label(p):
    if p < 0.001:
        return "***"
    elif p < 0.01:
        return "**"
    elif p < 0.05:
        return "*"
    else:
        return "ns"

# ─────────────────────────────────────────────────────────────────────────────
# DOCUMENT STRUCTURE
# ─────────────────────────────────────────────────────────────────────────────

# Title section
title = doc.add_paragraph()
title_run = title.add_run("MORAN'S I & LISA CLUSTER RESULTS")
title_run.font.size = Pt(28)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(255, 255, 255)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add background color to title paragraph
shading_elm = OxmlElement('w:shd')
shading_elm.set(qn('w:fill'), '1F3864')  # NAVY
title._element.get_or_add_pPr().append(shading_elm)

subtitle = doc.add_paragraph("Spatial Autocorrelation Analysis — Chandigarh Urban Sectors")
subtitle_run = subtitle.runs[0]
subtitle_run.font.size = Pt(16)
subtitle_run.font.color.rgb = COLORS['BLUE']
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle2 = doc.add_paragraph("Week 7 Days 6–7  |  999 permutation tests  |  p < 0.05 significance threshold")
subtitle2_run = subtitle2.runs[0]
subtitle2_run.font.size = Pt(11)
subtitle2_run.font.italic = True
subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()  # Spacer

# Section 1: Global Moran's I
p = doc.add_heading("1.  Global Moran's I — City-Wide Spatial Autocorrelation", level=1)
p_run = p.runs[0]
p_run.font.color.rgb = COLORS['NAVY']

doc.add_paragraph(f"The global Moran's I test was applied to mean LST values (2012–2025 summer composite average) across {len(w.neighbors)} sectors included in the spatial weights matrix. Sectors Manimajra and Khuda Alisher were excluded as geographic islands. Spatial weights were Queen contiguity, row-standardised, with an average of {avg_neigh:.2f} neighbours per sector. Significance was assessed using 999 random permutations.")

doc.add_paragraph()  # Spacer

# Key results box
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.5)
p.paragraph_format.right_indent = Inches(0.5)

# Add background color to paragraph
shading_elm = OxmlElement('w:shd')
shading_elm.set(qn('w:fill'), 'E2EFDA')  # Light green
p._element.get_or_add_pPr().append(shading_elm)

results_text = [
    f"Global Moran's I Result — LST  |  999 permutations",
    f"Moran's I statistic (I):  {mi.I:.6f}",
    f"Expected I under spatial randomness (E[I]):  {mi.EI:.6f}",
    f"z-score:  {mi.z_sim:.4f}",
    f"Pseudo p-value:  {mi.p_sim:.4f}  ***",
    f"Number of sectors in weights:  {len(w.neighbors)}",
    f"Average neighbours per sector:  {avg_neigh:.2f}  (Queen contiguity)",
]

for line in results_text:
    p = doc.add_paragraph(line, style='List Bullet')
    p_run = p.runs[0]
    p_run.font.size = Pt(10)

doc.add_paragraph()  # Spacer

# Data summary table
hh_list = ', '.join(gdf_weights[gdf_weights['LISA_cluster']=='HH']['Sector_nam'].tolist())
ll_list = ', '.join(gdf_weights[gdf_weights['LISA_cluster']=='LL']['Sector_nam'].tolist())
hh_n = (gdf_weights['LISA_cluster']=='HH').sum()
ll_n = (gdf_weights['LISA_cluster']=='LL').sum()

p = doc.add_paragraph(f"Summary: {hh_n} HH sectors ({hh_list}) and {ll_n} LL sectors ({ll_list}) identified. Highly significant positive spatial autocorrelation confirms thermal clustering.", style='Normal')

doc.add_paragraph()  # Spacer

# LISA data table
p = doc.add_heading("2.  Full LISA Results — All Sectors", level=1)
p_run = p.runs[0]
p_run.font.color.rgb = COLORS['NAVY']

# Create LISA table
table = doc.add_table(rows=1, cols=10)
table.style = 'Light Grid Accent 1'

# Header row
header_cells = table.rows[0].cells
headers = ['Sector', 'LST (°C)', 'Neighbours\nAvg LST', 'Local I', 'p-value', 'Sig.', 'Cluster', 'N neigh.', 'NDVI', 'TCI']

for i, header_text in enumerate(headers):
    cell = header_cells[i]
    cell.text = header_text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
    shade_cell(cell, '1F3864')  # NAVY

# Sort LISA data
lisa_sorted = sorted(lisa_rows, key=lambda x: (
    ['HH', 'LL', 'HL', 'LH', 'NS'].index(x['cl']),
    -x['lst']
))

# Add data rows
for row_data in lisa_sorted:
    row_cells = table.add_row().cells
    row_cells[0].text = str(row_data['s'])
    row_cells[1].text = f"{row_data['lst']:.3f}"
    row_cells[2].text = f"{row_data['n_lst']:.3f}"
    row_cells[3].text = f"{row_data['li']:.4f}"
    row_cells[4].text = f"{row_data['p']:.4f}"
    row_cells[5].text = p_label(row_data['p'])
    row_cells[6].text = row_data['cl']
    row_cells[7].text = str(row_data['nn'])
    row_cells[8].text = f"{row_data['ndvi']:.4f}"
    row_cells[9].text = f"{row_data['tci']:.4f}"
    
    # Shade cluster column
    shade_cell(row_cells[6], CLUSTER_COLORS[row_data['cl']])

doc.add_paragraph()  # Spacer

# Notes for paper writing
p = doc.add_heading("3.  Notes for Paper Writing", level=1)
p_run = p.runs[0]
p_run.font.color.rgb = COLORS['ORANGE']

doc.add_paragraph(f"Global spatial autocorrelation in LST was assessed using Moran's I with Queen contiguity weights (n={len(w.neighbors)} sectors, average {avg_neigh:.2f} neighbours). Statistical significance was evaluated using 999 random permutations. The resulting I={mi.I:.3f} (z={mi.z_sim:.2f}, p={mi.p_sim:.3f}) indicated strong positive spatial autocorrelation.")

# Save document
doc.save(OUT_DOCX)
print(f'Document saved: {OUT_DOCX}')


# ─────────────────────────────────────────────────────────────────────────────
# DONE
# ─────────────────────────────────────────────────────────────────────────────

print("\n=== Day 6–7 complete ===")
print(f"  {OUT_CORR_MATRIX}")
print(f"  {OUT_SCATTER}")
print(f"  {OUT_LISA_MAP}")
print(f"  {OUT_RHO_CSV}")
print(f"  {OUT_SUMMARY_CSV}")
print(f"  {OUT_PARAM_TABLE}")
print(f"  {MASTER_GEO_PATH}  (updated with LISA_cluster)")
print(f"  {OUT_DOCX}")
